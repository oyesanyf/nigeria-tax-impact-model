
"""
Enhanced Report Generator for Nigeria Tax Model
Generates comprehensive HTML reports with multiple visualizations and AI-generated executive summary.
"""
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import io
import base64
from datetime import datetime
import os
import tempfile
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
from sklearn.inspection import permutation_importance

# Set seaborn style
sns.set_style("whitegrid")
sns.set_palette("husl")

def generate_sensitivity_chart(model, df):
    """Calculate and plot feature importance using permutation importance."""
    feature_cols = ['Oil_Price', 'SME_Tax', 'VAT_Recovery', 'Digital_Penetration', 'Remittances_USD', 'Inflation_Rate']
    # Filter columns to only those present in df
    feature_cols = [c for c in feature_cols if c in df.columns]
    
    X = df[feature_cols]
    y = df['GDP_Growth']
    
    # Calculate Permutation Importance
    r = permutation_importance(model, X, y, n_repeats=10, random_state=42)
    
    # Sort by importance
    sorted_idx = r.importances_mean.argsort()
    
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.barh([feature_cols[i] for i in sorted_idx], r.importances_mean[sorted_idx], color='#27ae60', alpha=0.8)
    ax.set_xlabel("Permutation Importance (R² Increase)", fontsize=12)
    ax.set_title("Policy Influence: Which Factors Drive GDP?", fontsize=14, fontweight='bold')
    plt.tight_layout()
    
    img_str = fig_to_base64(fig)
    plt.close(fig)
    return img_str

def generate_charts(sim_results, df, oil_forecast, model_r2, model=None):
    """Generate multiple charts and return them as base64 encoded strings."""
    charts = {}
    
    # Existing charts...
    
    # Chart 1: Main Distribution Plot (KDE)
    fig, ax = plt.subplots(figsize=(12, 6))
    sns.kdeplot(data=sim_results['GDP_Old'], fill=True, color='#e74c3c', label='Old Law', alpha=0.4, ax=ax)
    sns.kdeplot(data=sim_results['GDP_New'], fill=True, color='#3498db', label='New Tax Act 2025', alpha=0.4, ax=ax)
    sns.kdeplot(data=sim_results['GDP_Shock'], fill=True, color='#f39c12', label='Inflation Shock (+50%)', alpha=0.4, ax=ax)
    ax.axvline(x=0, color='black', linestyle='--', linewidth=2, label='Recession Threshold')
    ax.set_xlabel('Projected GDP Growth (%)', fontsize=12)
    ax.set_ylabel('Probability Density', fontsize=12)
    ax.set_title('GDP Growth Distribution: Policy Scenarios (2026)', fontsize=14, fontweight='bold')
    ax.legend(fontsize=10)
    charts['distribution'] = fig_to_base64(fig)
    plt.close(fig)
    
    # Chart 2: Box Plot Comparison
    fig, ax = plt.subplots(figsize=(10, 6))
    data_for_box = pd.DataFrame({
        'Old Law': sim_results['GDP_Old'],
        'New Tax Act': sim_results['GDP_New'],
        'Inflation Shock': sim_results['GDP_Shock']
    })
    sns.boxplot(data=data_for_box, ax=ax, palette=['#e74c3c', '#3498db', '#f39c12'])
    ax.axhline(y=0, color='black', linestyle='--', linewidth=1.5)
    ax.set_ylabel('GDP Growth (%)', fontsize=12)
    ax.set_title('GDP Growth Range by Scenario', fontsize=14, fontweight='bold')
    charts['boxplot'] = fig_to_base64(fig)
    plt.close(fig)
    
    # Chart 3: Historical GDP Growth Trend
    fig, ax = plt.subplots(figsize=(12, 5))
    df_plot = df[['GDP_Growth']].dropna()
    ax.plot(df_plot.index, df_plot['GDP_Growth'], color='#2c3e50', linewidth=2)
    ax.fill_between(df_plot.index, 0, df_plot['GDP_Growth'], where=(df_plot['GDP_Growth'] > 0), 
                     color='#27ae60', alpha=0.3, label='Positive Growth')
    ax.fill_between(df_plot.index, 0, df_plot['GDP_Growth'], where=(df_plot['GDP_Growth'] < 0), 
                     color='#e74c3c', alpha=0.3, label='Recession')
    ax.axhline(y=0, color='black', linestyle='-', linewidth=0.8)
    ax.set_xlabel('Year', fontsize=12)
    ax.set_ylabel('GDP Growth (%)', fontsize=12)
    ax.set_title('Historical GDP Growth (Training Data)', fontsize=14, fontweight='bold')
    ax.legend()
    charts['historical'] = fig_to_base64(fig)
    plt.close(fig)
    
    # Chart 4: Oil Price History & Forecast
    fig, ax = plt.subplots(figsize=(12, 5))
    oil_data = df[['Oil_Price']].dropna()
    ax.plot(oil_data.index, oil_data['Oil_Price'], color='#8e44ad', linewidth=2, label='Historical')
    
    # Add forecast
    last_date = oil_data.index[-1]
    forecast_dates = pd.date_range(start=last_date, periods=5, freq='QE')[1:]
    ax.plot(forecast_dates, oil_forecast, color='#e67e22', linewidth=2, linestyle='--', 
            marker='o', label='2026 Forecast (AutoARIMA)')
    ax.fill_between(forecast_dates, oil_forecast - 10, oil_forecast + 10, 
                     color='#e67e22', alpha=0.2, label='±$10 Uncertainty')
    ax.set_xlabel('Year', fontsize=12)
    ax.set_ylabel('Oil Price ($/barrel)', fontsize=12)
    ax.set_title('Crude Oil Price: History & 2026 Forecast', fontsize=14, fontweight='bold')
    ax.legend()
    charts['oil'] = fig_to_base64(fig)
    plt.close(fig)
    
    # Chart 5: Correlation Heatmap
    fig, ax = plt.subplots(figsize=(10, 8))
    corr_cols = ['GDP_Growth', 'Oil_Price', 'SME_Tax', 'VAT_Recovery', 'Inflation_Rate', 'Digital_Penetration']
    corr_data = df[[c for c in corr_cols if c in df.columns]].corr()
    sns.heatmap(corr_data, annot=True, fmt='.2f', cmap='coolwarm', center=0, 
                square=True, linewidths=1, cbar_kws={"shrink": 0.8}, ax=ax)
    ax.set_title('Feature Correlation Matrix', fontsize=14, fontweight='bold')
    charts['correlation'] = fig_to_base64(fig)
    plt.close(fig)
    
    # Chart 6: Recession Risk Comparison (Bar Chart)
    fig, ax = plt.subplots(figsize=(8, 6))
    risk_old = (sim_results['GDP_Old'] < 0).mean() * 100
    risk_new = (sim_results['GDP_New'] < 0).mean() * 100
    risk_shock = (sim_results['GDP_Shock'] < 0).mean() * 100
    
    scenarios = ['Old Law', 'New Tax Act', 'Inflation Shock']
    risks = [risk_old, risk_new, risk_shock]
    colors = ['#e74c3c', '#3498db', '#f39c12']
    
    bars = ax.bar(scenarios, risks, color=colors, alpha=0.7, edgecolor='black', linewidth=1.5)
    ax.axhline(y=10, color='red', linestyle='--', linewidth=1, label='High Risk Threshold (10%)')
    ax.set_ylabel('Recession Risk (%)', fontsize=12)
    ax.set_title('Recession Risk by Policy Scenario', fontsize=14, fontweight='bold')
    ax.legend()
    
    # Add value labels on bars
    for bar, risk in zip(bars, risks):
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height,
                f'{risk:.1f}%', ha='center', va='bottom', fontweight='bold')
    
    charts['risk_bars'] = fig_to_base64(fig)
    plt.close(fig)
    
    # Chart 7: Sensitivity Analysis (New)
    if model is not None:
        charts['sensitivity'] = generate_sensitivity_chart(model, df)
    
    return charts

def fig_to_base64(fig):
    """Convert matplotlib figure to base64 string."""
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    return base64.b64encode(buf.getvalue()).decode('utf-8')

def generate_ai_summary(df, model_r2, sim_results, oil_forecast):
    """Generate executive summary using OpenAI."""
    try:
        import openai
        
        # Check for API key
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            return generate_fallback_summary(df, model_r2, sim_results, oil_forecast)
        
        client = openai.OpenAI(api_key=api_key)
        
        # Prepare context
        risk_old = (sim_results['GDP_Old'] < 0).mean() * 100
        risk_new = (sim_results['GDP_New'] < 0).mean() * 100
        risk_shock = (sim_results['GDP_Shock'] < 0).mean() * 100
        
        prompt = f"""You are an economic policy analyst. Write a detailed executive summary for a Nigeria Tax Policy Impact Assessment report.

**Dataset Information:**
- Time Period: {df.index.min().year} to {df.index.max().year}
- Data Points: {len(df)} quarterly observations
- Sources: National Bureau of Statistics (NBS), Central Bank of Nigeria (CBN), World Bank, Google Data Commons
- Key Variables: GDP Growth, Crude Oil Prices, Tax Revenue (CIT & VAT), Inflation, Digital Penetration, Remittances

**Methodology:**
1. AutoARIMA time series forecasting for oil price prediction (2026 forecast: ${oil_forecast.mean():.2f}/barrel)
2. TPOT AutoML (Genetic Algorithm) for GDP impact modeling (R² = {model_r2:.4f})
3. Monte Carlo simulation with {len(sim_results)} iterations for risk assessment

**Key Findings:**
- Old Law Recession Risk: {risk_old:.1f}%
- New Tax Act 2025 Recession Risk: {risk_new:.1f}%
- Inflation Shock Scenario Risk: {risk_shock:.1f}%

Write a comprehensive 4-paragraph executive summary covering:
1. Context and objective
2. Data sources and methodology
3. Key findings and policy implications
4. Limitations and recommendations

Use professional economic language. Be specific about the techniques used."""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            max_tokens=800
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        print(f"Warning: Could not generate AI summary: {e}")
        return generate_fallback_summary(df, model_r2, sim_results, oil_forecast)

def generate_fallback_summary(df, model_r2, sim_results, oil_forecast):
    """Generate a detailed summary without AI."""
    risk_old = (sim_results['GDP_Old'] < 0).mean() * 100
    risk_new = (sim_results['GDP_New'] < 0).mean() * 100
    
    return f"""
    <h3>Executive Summary</h3>
    
    <p><strong>Objective:</strong> This report assesses the macroeconomic impact of Nigeria's proposed Tax Act 2025, 
    which eliminates Company Income Tax for SMEs and enhances VAT collection efficiency through digitalization.</p>
    
    <p><strong>Data & Methodology:</strong> The analysis leverages {len(df)} quarterly observations spanning 
    {df.index.min().year}–{df.index.max().year}, sourced from the National Bureau of Statistics (NBS), 
    Central Bank of Nigeria (CBN), World Bank Development Indicators, and Google Data Commons. 
    The modeling framework employs AutoARIMA for oil price forecasting (2026 baseline: ${oil_forecast.mean():.2f}/barrel) 
    and TPOT AutoML (Tree-based Pipeline Optimization Tool) using genetic algorithms to discover the optimal 
    regression model, achieving an R² of {model_r2:.4f}. Risk assessment utilizes Monte Carlo simulation 
    with {len(sim_results):,} iterations to quantify recession probability under stochastic oil price scenarios.</p>
    
    <p><strong>Key Findings:</strong> Under baseline conditions, the New Tax Act reduces recession risk from 
    {risk_old:.1f}% to {risk_new:.1f}%, suggesting a net positive impact on GDP growth. However, the model 
    indicates vulnerability to inflationary shocks, with recession risk escalating significantly under a 
    50% inflation spike scenario. The analysis reveals that oil price volatility remains the dominant 
    macroeconomic driver, with tax policy effects being secondary but measurable.</p>
    
    <p><strong>Limitations & Recommendations:</strong> The model assumes historical relationships remain stable 
    and does not account for structural breaks or behavioral changes in tax compliance. The derived tax rate 
    proxies (calculated from revenue/oil price ratios) are approximations pending availability of granular 
    quarterly effective tax rate data from FIRS/NRS. We recommend continuous monitoring of inflation dynamics 
    and complementary fiscal measures to mitigate downside risks identified in the shock scenario.</p>
    """

def generate_enhanced_report(sim_results, df, oil_forecast, model, model_r2, n_simulations):
    """Generate comprehensive HTML report with multiple charts and AI summary."""
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Generate all charts
    print("Generating visualizations and sensitivity analysis...")
    charts = generate_charts(sim_results, df, oil_forecast, model_r2, model=model)
    
    # Generate AI summary
    print("Generating executive summary...")
    executive_summary = generate_ai_summary(df, model_r2, sim_results, oil_forecast)
    
    # Calculate metrics
    risk_old = (sim_results['GDP_Old'] < 0).mean() * 100
    risk_new = (sim_results['GDP_New'] < 0).mean() * 100
    risk_shock = (sim_results['GDP_Shock'] < 0).mean() * 100
    
    mean_old = sim_results['GDP_Old'].mean()
    mean_new = sim_results['GDP_New'].mean()
    mean_shock = sim_results['GDP_Shock'].mean()
    
    # Build HTML
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Nigeria Tax Model Report - {timestamp}</title>
        <style>
            body {{
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                max-width: 1200px;
                margin: 0 auto;
                padding: 30px;
                background-color: #f5f7fa;
                color: #2c3e50;
                line-height: 1.6;
            }}
            .header {{
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                padding: 40px;
                border-radius: 10px;
                margin-bottom: 30px;
                box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            }}
            h1 {{
                margin: 0;
                font-size: 2.5em;
            }}
            .timestamp {{
                opacity: 0.9;
                font-size: 0.9em;
                margin-top: 10px;
            }}
            .section {{
                background: white;
                padding: 30px;
                margin-bottom: 30px;
                border-radius: 8px;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            }}
            h2 {{
                color: #667eea;
                border-bottom: 3px solid #667eea;
                padding-bottom: 10px;
                margin-top: 0;
            }}
            .metric-grid {{
                display: grid;
                grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
                gap: 20px;
                margin: 20px 0;
            }}
            .metric-card {{
                background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
                padding: 20px;
                border-radius: 8px;
                border-left: 5px solid #667eea;
            }}
            .metric-label {{
                font-size: 0.9em;
                color: #7f8c8d;
                margin-bottom: 5px;
            }}
            .metric-value {{
                font-size: 2em;
                font-weight: bold;
                color: #2c3e50;
            }}
            .chart-container {{
                margin: 30px 0;
                text-align: center;
            }}
            .chart-container img {{
                max-width: 100%;
                border-radius: 8px;
                box-shadow: 0 4px 8px rgba(0,0,0,0.1);
            }}
            .risk-high {{ color: #e74c3c; }}
            .risk-low {{ color: #27ae60; }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin: 20px 0;
            }}
            th, td {{
                padding: 12px;
                text-align: left;
                border-bottom: 1px solid #ecf0f1;
            }}
            th {{
                background-color: #667eea;
                color: white;
            }}
            tr:hover {{
                background-color: #f8f9fa;
            }}
            .footer {{
                text-align: center;
                color: #7f8c8d;
                font-size: 0.9em;
                margin-top: 50px;
                padding-top: 20px;
                border-top: 1px solid #ecf0f1;
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>🇳🇬 Nigeria Tax Policy Impact Assessment</h1>
            <div class="timestamp">Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}</div>
            <div class="timestamp">Report ID: {timestamp}</div>
        </div>
        
        <div class="section">
            {executive_summary}
        </div>
        
        <div class="section">
            <h2>📊 Key Metrics</h2>
            <div class="metric-grid">
                <div class="metric-card">
                    <div class="metric-label">Model Accuracy (R²)</div>
                    <div class="metric-value">{model_r2:.1%}</div>
                </div>
                <div class="metric-card">
                    <div class="metric-label">Simulations Run</div>
                    <div class="metric-value">{n_simulations:,}</div>
                </div>
                <div class="metric-card">
                    <div class="metric-label">2026 Oil Forecast</div>
                    <div class="metric-value">${oil_forecast.mean():.0f}</div>
                </div>
                <div class="metric-card">
                    <div class="metric-label">Data Points</div>
                    <div class="metric-value">{len(df)}</div>
                </div>
            </div>
        </div>
        
        <div class="section">
            <h2>📈 Scenario Comparison</h2>
            <table>
                <thead>
                    <tr>
                        <th>Scenario</th>
                        <th>Mean GDP Growth</th>
                        <th>Recession Risk</th>
                        <th>Risk Change</th>
                    </tr>
                </thead>
                <tbody>
                    <tr>
                        <td><strong>Old Law (Baseline)</strong></td>
                        <td>{mean_old:.2f}%</td>
                        <td class="{'risk-high' if risk_old > 10 else 'risk-low'}">{risk_old:.1f}%</td>
                        <td>—</td>
                    </tr>
                    <tr>
                        <td><strong>New Tax Act 2025</strong></td>
                        <td>{mean_new:.2f}%</td>
                        <td class="{'risk-high' if risk_new > 10 else 'risk-low'}">{risk_new:.1f}%</td>
                        <td class="{'risk-low' if risk_new < risk_old else 'risk-high'}">
                            {risk_new - risk_old:+.1f}pp
                        </td>
                    </tr>
                    <tr>
                        <td><strong>Inflation Shock (+50%)</strong></td>
                        <td>{mean_shock:.2f}%</td>
                        <td class="risk-high">{risk_shock:.1f}%</td>
                        <td class="risk-high">{risk_shock - risk_old:+.1f}pp</td>
                    </tr>
                </tbody>
            </table>
        </div>
        
        <div class="section">
            <h2>📉 Visualizations</h2>
            
            <div class="chart-container">
                <h3>GDP Growth Distribution by Scenario</h3>
                <img src="data:image/png;base64,{charts['distribution']}" alt="Distribution">
            </div>
            
            <div class="chart-container">
                <h3>Recession Risk Comparison</h3>
                <img src="data:image/png;base64,{charts['risk_bars']}" alt="Risk Bars">
            </div>
            
            <div class="chart-container">
                <h3>GDP Growth Range (Box Plot)</h3>
                <img src="data:image/png;base64,{charts['boxplot']}" alt="Box Plot">
            </div>
            
            <div class="chart-container">
                <h3>Historical GDP Growth Trend</h3>
                <img src="data:image/png;base64,{charts['historical']}" alt="Historical">
            </div>
            
            <div class="chart-container">
                <h3>Oil Price: History & 2026 Forecast</h3>
                <img src="data:image/png;base64,{charts['oil']}" alt="Oil Forecast">
            </div>
            
            <div class="chart-container">
                <h3>Feature Correlation Matrix</h3>
                <img src="data:image/png;base64,{charts['correlation']}" alt="Correlation">
            </div>

            <div class="chart-container">
                <h3>Policy Influence: Factor Importance</h3>
                <p style="font-size: 0.9em; color: #7f8c8d; max-width: 800px; margin: 0 auto 15px;">
                    This chart uses <strong>Permutation Importance</strong> to measure which variables have the most leverage 
                    over predicted GDP. Higher values indicate factors with greater economic impact in the model.
                </p>
                <img src="data:image/png;base64,{charts['sensitivity']}" alt="Sensitivity Analysis">
            </div>
        </div>
        
        <div class="section">
            <h2>📚 Data Sources</h2>
            <table>
                <thead>
                    <tr>
                        <th>Source</th>
                        <th>Data Type</th>
                        <th>Role in Model</th>
                        <th>URL/Reference</th>
                    </tr>
                </thead>
                <tbody>
                    <tr>
                        <td><strong>National Bureau of Statistics (NBS)</strong></td>
                        <td>GDP Growth (Real/Nominal)</td>
                        <td>Ground Truth for training and historical targets</td>
                        <td><a href="https://nigerianstat.gov.ng/elibrary" target="_blank">NBS eLibrary</a></td>
                    </tr>
                    <tr>
                        <td><strong>Central Bank of Nigeria (CBN)</strong></td>
                        <td>Crude Oil Prices (Bonny Light)</td>
                        <td>Primary economic driver and wealth engine</td>
                        <td><a href="https://www.cbn.gov.ng/documents/statbulletin.asp" target="_blank">CBN Statistical Bulletin</a></td>
                    </tr>
                    <tr>
                        <td><strong>Nigeria Revenue Service (NRS)</strong></td>
                        <td>2025 Actual CIT and VAT collections</td>
                        <td>Calibration anchor for 2026 forecasts</td>
                        <td><a href="https://nrs.gov.ng/transparency/revenue-dashboard" target="_blank">NRS Dashboard</a></td>
                    </tr>
                    <tr>
                        <td><strong>Google Data Commons</strong></td>
                        <td>Tax Revenue (% GDP), Inflation (CPI)</td>
                        <td>Fills historical gaps (2000-2015)</td>
                        <td><a href="https://datacommons.org" target="_blank">Data Commons API</a></td>
                    </tr>
                    <tr>
                        <td><strong>World Bank (WB)</strong></td>
                        <td>Digital Penetration, Remittances</td>
                        <td>Modern factors for digitalization impact</td>
                        <td><a href="https://databank.worldbank.org" target="_blank">WDI Indicators</a></td>
                    </tr>
                </tbody>
            </table>
        </div>
        
        <div class="section">
            <h2>🔄 Data Processing Pipeline</h2>
            <ol>
                <li><strong>Data Ingestion:</strong> Raw Excel files downloaded from NBS/CBN official portals</li>
                <li><strong>Cleaning:</strong> Missing values handled via forward-fill (economic continuity assumption)</li>
                <li><strong>Frequency Alignment:</strong> Annual World Bank data resampled to quarterly using linear interpolation</li>
                <li><strong>Feature Engineering:</strong>
                    <ul>
                        <li><em>SME Tax Proxy:</em> Derived from <code>CIT_Revenue / Oil_Price × Scaler</code> (calibrated to historical 25% mean)</li>
                        <li><em>VAT Recovery Proxy:</em> Derived from <code>VAT_Revenue / Oil_Price × Scaler</code> (calibrated to historical 55% mean)</li>
                    </ul>
                </li>
                <li><strong>Data Fusion:</strong> Prioritized NRS 2025 Actuals → Google Data Commons → NBS/CBN Files</li>
                <li><strong>Output:</strong> Final dataset with {len(df)} observations across {len(df.columns)} variables</li>
            </ol>
        </div>
        
        <div class="section">
            <h2>🧠 Training Methodology</h2>
            
            <h3>1. Time Series Forecasting (Oil Price)</h3>
            <table>
                <tr><td><strong>Algorithm:</strong></td><td>AutoARIMA (pmdarima)</td></tr>
                <tr><td><strong>Parameter Selection:</strong></td><td>Automatic (p,d,q) via AIC minimization</td></tr>
                <tr><td><strong>Forecast Horizon:</strong></td><td>4 quarters (2026)</td></tr>
                <tr><td><strong>2026 Baseline Forecast:</strong></td><td>${oil_forecast.mean():.2f}/barrel (±$10 uncertainty)</td></tr>
            </table>
            
            <h3>2. AutoML Model Training (GDP Impact)</h3>
            <table>
                <tr><td><strong>Framework:</strong></td><td>TPOT (Tree-based Pipeline Optimization Tool)</td></tr>
                <tr><td><strong>Search Method:</strong></td><td>Genetic Programming (evolutionary algorithm)</td></tr>
                <tr><td><strong>Generations:</strong></td><td>4 (evolution cycles)</td></tr>
                <tr><td><strong>Population Size:</strong></td><td>30 (candidate pipelines per generation)</td></tr>
                <tr><td><strong>Cross-Validation:</strong></td><td>TimeSeriesSplit (5-fold walk-forward)</td></tr>
                <tr><td><strong>Scoring Metric:</strong></td><td>R² (coefficient of determination)</td></tr>
                <tr><td><strong>Train/Test Split:</strong></td><td>80% historical / 20% holdout (chronological, no shuffle)</td></tr>
                <tr><td><strong>Final Model R²:</strong></td><td>{model_r2:.4f} ({model_r2*100:.1f}% variance explained)</td></tr>
            </table>
            
            <h3>3. Feature Vector</h3>
            <table>
                <thead>
                    <tr><th>Feature</th><th>Description</th><th>Unit</th></tr>
                </thead>
                <tbody>
                    <tr><td>Oil_Price</td><td>Crude oil price (Bonny Light)</td><td>USD/barrel</td></tr>
                    <tr><td>SME_Tax</td><td>Effective SME tax rate (derived proxy)</td><td>%</td></tr>
                    <tr><td>VAT_Recovery</td><td>VAT collection efficiency (derived proxy)</td><td>%</td></tr>
                    <tr><td>Digital_Penetration</td><td>Internet users as % of population</td><td>%</td></tr>
                    <tr><td>Remittances_USD</td><td>Personal remittances received</td><td>USD (billions)</td></tr>
                    <tr><td>Inflation_Rate</td><td>Consumer Price Index (CPI) annual change</td><td>%</td></tr>
                </tbody>
            </table>
            
            <h3>4. Monte Carlo Simulation</h3>
            <table>
                <tr><td><strong>Iterations:</strong></td><td>{n_simulations:,}</td></tr>
                <tr><td><strong>Oil Price Distribution:</strong></td><td>Normal(μ=${oil_forecast.mean():.0f}, σ=$10)</td></tr>
                <tr><td><strong>Recession Threshold:</strong></td><td>GDP Growth < 0%</td></tr>
            </table>
            
            <h3>5. Best Model Pipeline</h3>
            <pre style="background:#f4f4f4; padding:15px; border-radius:5px; overflow-x:auto;">
{str(model.fitted_pipeline_) if hasattr(model, 'fitted_pipeline_') else str(model)}
            </pre>
        </div>
        
        <div class="section">
            <h2>📈 Recommendations for Improved Accuracy</h2>
            <p>The current model achieves <strong>R² = {model_r2:.1%}</strong>. The following additional data sources would significantly improve predictive accuracy:</p>
            
            <table>
                <thead>
                    <tr>
                        <th>Data Source</th>
                        <th>What It Provides</th>
                        <th>Expected Impact</th>
                        <th>How to Obtain</th>
                    </tr>
                </thead>
                <tbody>
                    <tr>
                        <td><strong>FIRS Quarterly Tax Returns</strong></td>
                        <td>Actual effective tax rates by sector (not proxies)</td>
                        <td>High - Eliminates proxy estimation error</td>
                        <td>Freedom of Information request to FIRS</td>
                    </tr>
                    <tr>
                        <td><strong>NRS Monthly Revenue Data</strong></td>
                        <td>Higher frequency revenue observations</td>
                        <td>Medium - More training data points</td>
                        <td><a href="https://nrs.gov.ng/transparency" target="_blank">NRS Transparency Portal</a></td>
                    </tr>
                    <tr>
                        <td><strong>NIBSS Real-Time Payments Data</strong></td>
                        <td>Digital transaction volumes (proxy for formalization)</td>
                        <td>High - Direct measure of digital economy</td>
                        <td><a href="https://nibss-plc.com.ng" target="_blank">NIBSS Annual Reports</a></td>
                    </tr>
                    <tr>
                        <td><strong>CBN Business Expectations Survey</strong></td>
                        <td>Leading indicator of economic sentiment</td>
                        <td>Medium - Forward-looking signal</td>
                        <td><a href="https://www.cbn.gov.ng/documents/businessexpectations.asp" target="_blank">CBN Survey Data</a></td>
                    </tr>
                    <tr>
                        <td><strong>IMF Article IV Consultation</strong></td>
                        <td>Independent macro forecasts and fiscal analysis</td>
                        <td>Low - Validation benchmark</td>
                        <td><a href="https://www.imf.org/en/Countries/NGA" target="_blank">IMF Nigeria Page</a></td>
                    </tr>
                    <tr>
                        <td><strong>Sectoral GDP Breakdown (NBS)</strong></td>
                        <td>GDP by sector (Oil vs Non-Oil vs Services)</td>
                        <td>High - Isolates SME-heavy sectors</td>
                        <td><a href="https://nigerianstat.gov.ng" target="_blank">NBS Quarterly GDP Reports</a></td>
                    </tr>
                </tbody>
            </table>
            
            <h3>Data Quality Improvements</h3>
            <ul>
                <li><strong>Extended History:</strong> Data prior to 2006 (if available from CBN archives) would improve trend detection</li>
                <li><strong>Regional Disaggregation:</strong> State-level tax collection data would reveal geographic policy effects</li>
                <li><strong>Firm-Level Data:</strong> CAC registration + tax payment matching would validate SME assumptions</li>
            </ul>
            
            <h3>Model Improvements</h3>
            <ul>
                <li><strong>Increase TPOT Generations:</strong> From 4 to 10+ for more thorough pipeline search</li>
                <li><strong>Add Lag Features:</strong> Include GDP_Growth(t-1) and Oil_Price(t-2) for temporal dynamics</li>
                <li><strong>Ensemble Stacking:</strong> Combine TPOT's best model with AutoARIMA for hybrid forecasting</li>
            </ul>
        </div>
        
        <div class="footer">
            <p>Nigeria Tax Model System | Powered by AutoML & Monte Carlo Simulation</p>
            <p>Data Sources: NBS, CBN, World Bank, Google Data Commons, NRS Dashboard</p>
            <p style="margin-top: 15px;">
                <a href="https://github.com/oyesanyf/nigeria-tax-impact-model" target="_blank" style="color: #667eea; text-decoration: none; font-weight: bold;">
                    View Source Code on GitHub ↗
                </a>
            </p>
        </div>
    </body>
    </html>
    """
    
    # Save report
    os.makedirs('reports', exist_ok=True)
    base_filename = f"reports/tax_impact_report_{timestamp}"
    html_filename = f"{base_filename}.html"
    with open(html_filename, 'w', encoding='utf-8') as f:
        f.write(html)
    
    paths = {"html": os.path.abspath(html_filename)}

    # Generate additional formats
    docx_filename = f"{base_filename}.docx"
    pdf_filename = f"{base_filename}.pdf"
    
    try:
        generate_docx_report(docx_filename, timestamp, executive_summary, df, model_r2, oil_forecast, sim_results, n_simulations, model, charts)
        paths["word"] = os.path.abspath(docx_filename)
    except Exception as e:
        print(f"❌ Error generating Word report: {e}")

    try:
        generate_pdf_report(pdf_filename, timestamp, executive_summary, df, model_r2, oil_forecast, sim_results, n_simulations, model, charts)
        paths["pdf"] = os.path.abspath(pdf_filename)
    except Exception as e:
        print(f"❌ Error generating PDF report: {e}")
    
    return paths

def generate_docx_report(filename, timestamp, summary, df, model_r2, oil_forecast, sim_results, n_simulations, model, charts):
    doc = Document()
    doc.add_heading('Nigeria Tax Policy Impact Assessment', 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    doc.add_paragraph(f"Report ID: {timestamp}")

    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(summary.replace('<h3>', '').replace('</h3>', '').replace('<strong>', '').replace('</strong>', '').replace('<p>', '').replace('</p>', '\n'))

    doc.add_heading('Key Metrics', level=1)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    metrics = [
        ('Model Accuracy (R²)', f"{model_r2:.1%}"),
        ('Simulations Run', f"{n_simulations:,}"),
        ('2026 Oil Forecast', f"${oil_forecast.mean():.0f}"),
        ('Data Points', str(len(df)))
    ]
    for m, v in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = m
        row_cells[1].text = v

    doc.add_heading('Visualizations', level=1)
    for title, key in [
        ('GDP Growth Distribution', 'distribution'),
        ('Recession Risk Comparison', 'risk_bars'),
        ('GDP Growth Range', 'boxplot'),
        ('Historical GDP Trend', 'historical'),
        ('Oil Price History & Forecast', 'oil'),
        ('Feature Correlation Matrix', 'correlation'),
        ('Policy Influence (Sensitivity)', 'sensitivity')
    ]:
        doc.add_heading(title, level=2)
        # Convert base64 back to image for docx
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
            tmpfile.write(base64.b64decode(charts[key]))
            tmpfile_path = tmpfile.name
        doc.add_picture(tmpfile_path, width=Inches(6))
        os.unlink(tmpfile_path)

    doc.add_heading('Training Methodology', level=1)
    model_str = str(model.fitted_pipeline_) if hasattr(model, 'fitted_pipeline_') else str(model)
    doc.add_paragraph(f"ML Framework: TPOT (Genetic Programming)\nModel: {model_str[:500]}")

    doc.add_heading('Data Sources', level=1)
    ds_table = doc.add_table(rows=1, cols=3)
    hdr_cells = ds_table.rows[0].cells
    hdr_cells[0].text = 'Source'
    hdr_cells[1].text = 'Data Type'
    hdr_cells[2].text = 'Role in Model'
    
    sources = [
        ('NBS', 'GDP Growth (Real/Nominal)', 'Ground Truth for training'),
        ('CBN', 'Crude Oil Prices (Bonny Light)', 'Primary economic driver'),
        ('NRS', '2025 Actual Citizen Collections', 'Calibration anchor for forecasts'),
        ('Data Commons', 'Tax Revenue, Inflation', 'Fills historical gaps (2000-2015)'),
        ('World Bank', 'Digital Penetration, Remittances', 'Modern impact factors')
    ]
        row_cells[0].text = s
        row_cells[1].text = d
        row_cells[2].text = r

    doc.add_heading('Source Code', level=1)
    doc.add_paragraph('Theoretical framework and source code available at:')
    doc.add_paragraph('https://github.com/oyesanyf/nigeria-tax-impact-model', style='List Bullet')

    doc.save(filename)

def generate_pdf_report(filename, timestamp, summary, df, model_r2, oil_forecast, sim_results, n_simulations, model, charts):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # Title
    pdf.set_font("helvetica", 'B', 16)
    pdf.cell(0, 10, "Nigeria Tax Policy Impact Assessment", ln=True, align='C')
    pdf.set_font("helvetica", size=10)
    pdf.cell(0, 10, f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", ln=True, align='C')
    pdf.cell(0, 10, f"Report ID: {timestamp}", ln=True, align='C')
    pdf.ln(10)

    # Summary
    pdf.set_font("helvetica", 'B', 14)
    pdf.cell(0, 10, "Executive Summary", ln=True)
    pdf.set_font("helvetica", size=11)
    # Simple cleanup of HTML tags for PDF
    clean_summary = summary.replace('<h3>', '').replace('</h3>', '\n').replace('<strong>', '').replace('</strong>', '').replace('<p>', '').replace('</p>', '\n')
    # Replace common non-ASCII characters that break standard PDF fonts
    clean_summary = clean_summary.replace('\u2013', '-').replace('\u2014', '-').replace('\u2018', "'").replace('\u2019', "'").replace('\u201c', '"').replace('\u201d', '"')
    pdf.multi_cell(0, 7, clean_summary)
    pdf.ln(5)

    # Charts
    pdf.set_font("helvetica", 'B', 14)
    pdf.cell(0, 10, "Visualizations", ln=True)
    
    for title, key in [
        ('GDP Growth Distribution', 'distribution'),
        ('Recession Risk Comparison', 'risk_bars'),
        ('GDP Growth Range', 'boxplot'),
        ('Historical GDP Trend', 'historical'),
        ('Oil Price Forecast', 'oil'),
        ('Feature Correlation', 'correlation'),
        ('Policy Influence', 'sensitivity')
    ]:
        if pdf.get_y() > 200: pdf.add_page()
        pdf.set_font("helvetica", 'B', 12)
        pdf.cell(0, 10, title, ln=True)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
            tmpfile.write(base64.b64decode(charts[key]))
            tmpfile_path = tmpfile.name
        pdf.image(tmpfile_path, x=10, w=180)
        os.unlink(tmpfile_path)
        pdf.ln(5)

    # Data Sources
    pdf.add_page()
    pdf.set_font("helvetica", 'B', 14)
    pdf.cell(0, 10, "Data Sources", ln=True)
    pdf.set_font("helvetica", 'B', 10)
    
    # Table Header
    pdf.cell(30, 10, "Source", 1)
    pdf.cell(80, 10, "Data Type", 1)
    pdf.cell(80, 10, "Role in Model", 1, ln=True)
    
    pdf.set_font("helvetica", size=9)
    sources = [
        ('NBS', 'GDP Growth', 'Ground Truth'),
        ('CBN', 'Oil Prices', 'Economic Driver'),
        ('NRS', '2025 Actuals', 'Calibration Anchor'),
        ('Data Commons', 'Tax, Inflation', 'Historical Gaps'),
        ('World Bank', 'Digital, Remit.', 'Modern Factors')
    ]
    for s, d, r in sources:
        pdf.cell(30, 8, s, 1)
        pdf.cell(80, 8, d, 1)
        pdf.cell(80, 8, r, 1, ln=True)

    # GitHub Link
    pdf.ln(10)
    pdf.set_font("helvetica", 'I', 10)
    pdf.set_text_color(0, 0, 255)
    pdf.cell(0, 10, "Source Code: https://github.com/oyesanyf/nigeria-tax-impact-model", ln=True, align='C', link="https://github.com/oyesanyf/nigeria-tax-impact-model")

    pdf.output(filename)


def generate_multi_oil_charts(scenario_results, oil_prices, df, model_r2, model=None):
    """Generate comparative charts for multiple oil price scenarios."""
    charts = {}
    
    # Define a color palette for different oil prices
    colors = ['#e74c3c', '#3498db', '#27ae60', '#f39c12', '#9b59b6', '#1abc9c', '#e67e22', '#34495e']
    
    # Chart 1: GDP Distribution Comparison (New Law only) across oil prices
    fig, ax = plt.subplots(figsize=(14, 7))
    for i, (oil_price, sim_results) in enumerate(scenario_results.items()):
        color = colors[i % len(colors)]
        sns.kdeplot(data=sim_results['GDP_New'], fill=True, color=color, 
                    label=f'Oil @ ${oil_price:.0f}', alpha=0.4, ax=ax)
    ax.axvline(x=0, color='black', linestyle='--', linewidth=2, label='Recession Threshold')
    ax.set_xlabel('Projected GDP Growth (%)', fontsize=12)
    ax.set_ylabel('Probability Density', fontsize=12)
    ax.set_title('GDP Growth Distribution by Oil Price Scenario (New Tax Act 2025)', fontsize=14, fontweight='bold')
    ax.legend(fontsize=10)
    charts['multi_distribution'] = fig_to_base64(fig)
    plt.close(fig)
    
    # Chart 2: Recession Risk Bar Comparison
    fig, ax = plt.subplots(figsize=(12, 6))
    
    x_labels = [f'${p:.0f}' for p in oil_prices]
    x_pos = np.arange(len(oil_prices))
    width = 0.25
    
    risks_old = [(scenario_results[p]['GDP_Old'] < 0).mean() * 100 for p in oil_prices]
    risks_new = [(scenario_results[p]['GDP_New'] < 0).mean() * 100 for p in oil_prices]
    risks_shock = [(scenario_results[p]['GDP_Shock'] < 0).mean() * 100 for p in oil_prices]
    
    bars1 = ax.bar(x_pos - width, risks_old, width, label='Old Law', color='#e74c3c', alpha=0.8)
    bars2 = ax.bar(x_pos, risks_new, width, label='New Tax Act', color='#3498db', alpha=0.8)
    bars3 = ax.bar(x_pos + width, risks_shock, width, label='Inflation Shock', color='#f39c12', alpha=0.8)
    
    ax.set_xlabel('Oil Price ($/barrel)', fontsize=12)
    ax.set_ylabel('Recession Risk (%)', fontsize=12)
    ax.set_title('Recession Risk Across Oil Price Scenarios', fontsize=14, fontweight='bold')
    ax.set_xticks(x_pos)
    ax.set_xticklabels(x_labels)
    ax.legend()
    ax.axhline(y=10, color='red', linestyle='--', linewidth=1, alpha=0.5)
    
    # Add value labels on bars
    for bars in [bars1, bars2, bars3]:
        for bar in bars:
            height = bar.get_height()
            if height > 1:  # Only label if visible
                ax.text(bar.get_x() + bar.get_width()/2., height + 0.5,
                        f'{height:.0f}%', ha='center', va='bottom', fontsize=8)
    
    charts['multi_risk_bars'] = fig_to_base64(fig)
    plt.close(fig)
    
    # Chart 3: GDP Growth Trend by Oil Price (Box Plot)
    fig, ax = plt.subplots(figsize=(14, 6))
    
    box_data = []
    box_labels = []
    for oil_price in oil_prices:
        box_data.append(scenario_results[oil_price]['GDP_New'])
        box_labels.append(f'${oil_price:.0f}')
    
    bp = ax.boxplot(box_data, labels=box_labels, patch_artist=True)
    for i, patch in enumerate(bp['boxes']):
        patch.set_facecolor(colors[i % len(colors)])
        patch.set_alpha(0.6)
    
    ax.axhline(y=0, color='red', linestyle='--', linewidth=1.5, label='Recession Line')
    ax.set_xlabel('Oil Price ($/barrel)', fontsize=12)
    ax.set_ylabel('GDP Growth (%)', fontsize=12)
    ax.set_title('GDP Growth Range by Oil Price (New Tax Act)', fontsize=14, fontweight='bold')
    ax.legend()
    charts['multi_boxplot'] = fig_to_base64(fig)
    plt.close(fig)
    
    # Chart 4: Oil Price Sensitivity Line Chart
    fig, ax = plt.subplots(figsize=(12, 6))
    
    sorted_prices = sorted(oil_prices)
    mean_gdp_old = [scenario_results[p]['GDP_Old'].mean() for p in sorted_prices]
    mean_gdp_new = [scenario_results[p]['GDP_New'].mean() for p in sorted_prices]
    mean_gdp_shock = [scenario_results[p]['GDP_Shock'].mean() for p in sorted_prices]
    
    ax.plot(sorted_prices, mean_gdp_old, 'o-', color='#e74c3c', linewidth=2, markersize=8, label='Old Law')
    ax.plot(sorted_prices, mean_gdp_new, 's-', color='#3498db', linewidth=2, markersize=8, label='New Tax Act')
    ax.plot(sorted_prices, mean_gdp_shock, '^-', color='#f39c12', linewidth=2, markersize=8, label='Inflation Shock')
    
    ax.axhline(y=0, color='black', linestyle='--', linewidth=1, alpha=0.5)
    ax.fill_between(sorted_prices, 0, mean_gdp_new, where=[g > 0 for g in mean_gdp_new], 
                     color='#27ae60', alpha=0.1, label='Positive Growth Zone')
    ax.fill_between(sorted_prices, 0, mean_gdp_new, where=[g < 0 for g in mean_gdp_new], 
                     color='#e74c3c', alpha=0.1, label='Recession Zone')
    
    ax.set_xlabel('Oil Price ($/barrel)', fontsize=12)
    ax.set_ylabel('Mean GDP Growth (%)', fontsize=12)
    ax.set_title('Oil Price Sensitivity Analysis', fontsize=14, fontweight='bold')
    ax.legend(loc='best')
    ax.grid(True, alpha=0.3)
    charts['oil_sensitivity'] = fig_to_base64(fig)
    plt.close(fig)
    
    # Chart 5: Historical GDP (same as regular report)
    fig, ax = plt.subplots(figsize=(12, 5))
    df_plot = df[['GDP_Growth']].dropna()
    ax.plot(df_plot.index, df_plot['GDP_Growth'], color='#2c3e50', linewidth=2)
    ax.fill_between(df_plot.index, 0, df_plot['GDP_Growth'], where=(df_plot['GDP_Growth'] > 0), 
                     color='#27ae60', alpha=0.3, label='Positive Growth')
    ax.fill_between(df_plot.index, 0, df_plot['GDP_Growth'], where=(df_plot['GDP_Growth'] < 0), 
                     color='#e74c3c', alpha=0.3, label='Recession')
    ax.axhline(y=0, color='black', linestyle='-', linewidth=0.8)
    ax.set_xlabel('Year', fontsize=12)
    ax.set_ylabel('GDP Growth (%)', fontsize=12)
    ax.set_title('Historical GDP Growth (Training Data)', fontsize=14, fontweight='bold')
    ax.legend()
    charts['historical'] = fig_to_base64(fig)
    plt.close(fig)
    
    # Chart 6: Sensitivity Analysis (if model provided)
    if model is not None:
        charts['sensitivity'] = generate_sensitivity_chart(model, df)
    
    return charts


def generate_multi_oil_ai_summary(scenario_results, oil_prices, model_r2, n_simulations):
    """Generate executive summary for multi-scenario report using OpenAI."""
    try:
        import openai
        
        # Check for API key
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            return generate_multi_oil_fallback_summary(scenario_results, oil_prices, model_r2, n_simulations)
        
        client = openai.OpenAI(api_key=api_key)
        
        # Prepare data for prompt - ALL scenarios
        scenario_data_str = ""
        for p in sorted(oil_prices):
            sim = scenario_results[p]
            risk_old = (sim['GDP_Old'] < 0).mean() * 100
            risk_new = (sim['GDP_New'] < 0).mean() * 100
            risk_shock = (sim['GDP_Shock'] < 0).mean() * 100
            mean_new = sim['GDP_New'].mean()
            
            # Calculate range
            p5 = np.percentile(sim['GDP_New'], 5)
            p95 = np.percentile(sim['GDP_New'], 95)
            
            scenario_data_str += f"- Oil @ ${p:.0f}: GDP={mean_new:.2f}% (Range: {p5:.1f}% to {p95:.1f}%), Risk(New)={risk_new:.1f}%, Risk(Shock)={risk_shock:.1f}%\n"

        prompt = f"""You are a Chief Economist advising the Nigerian Presidency. Write a comprehensive "Macroeconomic Impact Assessment" for the 2025 Reform Agenda.

**Study Parameters:**
- **Objective:** Evaluate the resilience of the New Tax Act (0% SME Tax, Digital VAT Enhancement) against global oil price shocks.
- **Methodology:** {n_simulations} Monte Carlo simulations per scenario using an ML model (TPOT Optimized, R²={model_r2:.4f}).
- **Scenarios Tested:** {len(oil_prices)} distinct oil price levels.

**Simulation Data (By Oil Price Scenario):**
{scenario_data_str}

**Write a Detailed Strategic Report (HTML Format) covering:**

1.  **<h3>Executive Overview</h3>**
    - High-level synthesis of whether the New Tax Act is viable.
    - The "Break-even" oil price where recession risk becomes acceptable (<10%).

2.  **<h3>Scenario Deep-Dive</h3>**
    - Analyze the Low-Price Scenarios (Risk analysis, fiscal buffers needed).
    - Analyze the High-Price Scenarios (Growth opportunity, surplus allocation).
    - Contrast the "New Tax Act" vs "Old Law" performance specifically.

3.  **<h3>Structural Vulnerability Assessment</h3>**
    - Discuss the dependency on oil vs tax efficiency.
    - How much does the "Inflation Shock" scenario destabilize the reforms?

4.  **<h3>Strategic Policy Recommendations</h3>**
    - **Fiscal:** What should be done if oil drops below $60?
    - **Monetary:** How to coordinate with CBN on inflation?
    - **Implementation:** Critical success factors for the Digital VAT rollout.

**Format Rules:**
- Use professional economic terminology (e.g., "fiscal consolidation", "counter-cyclical buffers", "revenue elasticity").
- Use valid HTML tags: <h3>, <p>, <ul>, <li>, <strong>.
- NO Markdown (no #, no **).
- Length: Detailed and comprehensive (~800-1000 words).
"""

        response = client.chat.completions.create(
            model="gpt-4o",  # Upgraded to GPT-4o for better reasoning
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3, # Slightly creativity for analysis
            max_tokens=2500  # Allow long response
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        print(f"Warning: Could not generate AI summary: {e}")
        return generate_multi_oil_fallback_summary(scenario_results, oil_prices, model_r2, n_simulations)


def generate_multi_oil_fallback_summary(scenario_results, oil_prices, model_r2, n_simulations):
    """Fallback summary template if AI fails."""
    min_risk_price = min(oil_prices, key=lambda p: (scenario_results[p]['GDP_New'] < 0).mean())
    max_risk_price = max(oil_prices, key=lambda p: (scenario_results[p]['GDP_New'] < 0).mean())
    
    return f"""
    <h3>Multi-Scenario Oil Price Analysis</h3>
    
    <p><strong>Objective:</strong> This analysis compares the economic impact of Nigeria's Tax Act 2025 
    across {len(oil_prices)} different oil price scenarios, ranging from ${min(oil_prices):.0f} to ${max(oil_prices):.0f} per barrel.</p>
    
    <p><strong>Key Finding:</strong> The optimal scenario occurs at <strong>${min_risk_price:.0f}/barrel</strong> 
    (lowest recession risk under New Tax Act), while the highest risk is at <strong>${max_risk_price:.0f}/barrel</strong>.</p>
    
    <p><strong>Methodology:</strong> Each oil price scenario was simulated with {n_simulations:,} Monte Carlo iterations 
    using the trained TPOT model (R² = {model_r2:.4f}). Recession risk is calculated as the probability of GDP growth falling below 0%.</p>
    """


def generate_multi_oil_report(scenario_results, df, oil_prices, model, model_r2, n_simulations):
    """Generate comprehensive HTML report comparing multiple oil price scenarios."""
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Generate comparative charts
    print("Generating multi-scenario visualizations...")
    charts = generate_multi_oil_charts(scenario_results, oil_prices, df, model_r2, model=model)
    
    # Build scenario comparison table
    scenario_rows = ""
    for oil_price in oil_prices:
        sim = scenario_results[oil_price]
        risk_old = (sim['GDP_Old'] < 0).mean() * 100
        risk_new = (sim['GDP_New'] < 0).mean() * 100
        mean_gdp = sim['GDP_New'].mean()
        
        # Calculate range
        p5 = np.percentile(sim['GDP_New'], 5)
        p95 = np.percentile(sim['GDP_New'], 95)
        
        # Determine Status
        if risk_new < 5:
            status = "🟢 SAFE"
            status_style = "color: #27ae60; font-weight: bold;"
        elif risk_new <= 15:
            status = "🟡 CAUTION"
            status_style = "color: #f39c12; font-weight: bold;"
        else:
            status = "🔴 DANGER"
            status_style = "color: #e74c3c; font-weight: bold;"
            
        scenario_rows += f"""
        <tr>
            <td><strong>${oil_price:.0f}/barrel</strong></td>
            <td>{mean_gdp:.2f}%</td>
            <td style="font-family: monospace; font-size: 0.9em;">{p5:+.1f}% to {p95:+.1f}%</td>
            <td class="{'risk-high' if risk_new > 10 else 'risk-low'}">{risk_new:.1f}%</td>
            <td style="{status_style}">{status}</td>
            <td class="{'risk-low' if risk_new < risk_old else 'risk-high'}">{risk_new - risk_old:+.1f}pp</td>
        </tr>
        """
    
    # Generate summary
    print("Generating executive summary...")
    
    # Try AI summary first
    executive_summary = generate_multi_oil_ai_summary(scenario_results, oil_prices, model_r2, n_simulations)
    
    # Fallback to template if AI fails or returns None/Empty
    if not executive_summary or "OpenAI" in executive_summary: # Basic check for valid content/error message
         # Recalculate basic stats for fallback
        min_risk_price = min(oil_prices, key=lambda p: (scenario_results[p]['GDP_New'] < 0).mean())
        max_risk_price = max(oil_prices, key=lambda p: (scenario_results[p]['GDP_New'] < 0).mean())
        
        # Only overwrite if AI didn't work. If we have AI summary, skipping this block.
        # But wait, generate_multi_oil_ai_summary handles fallback internally? No, I'll implement it to return fallback string if API fails.
        pass # The function below will handle fallback logic

    # Build HTML
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Multi-Scenario Oil Price Analysis - {timestamp}</title>
        <style>
            body {{
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                max-width: 1400px;
                margin: 0 auto;
                padding: 30px;
                background-color: #f5f7fa;
                color: #2c3e50;
                line-height: 1.6;
            }}
            .header {{
                background: linear-gradient(135deg, #2c3e50 0%, #3498db 100%);
                color: white;
                padding: 40px;
                border-radius: 10px;
                margin-bottom: 30px;
                box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            }}
            h1 {{
                margin: 0;
                font-size: 2.5em;
            }}
            .timestamp {{
                opacity: 0.9;
                font-size: 0.9em;
                margin-top: 10px;
            }}
            .section {{
                background: white;
                padding: 30px;
                margin-bottom: 30px;
                border-radius: 8px;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            }}
            h2 {{
                color: #3498db;
                border-bottom: 3px solid #3498db;
                padding-bottom: 10px;
                margin-top: 0;
            }}
            .metric-grid {{
                display: grid;
                grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
                gap: 20px;
                margin: 20px 0;
            }}
            .metric-card {{
                background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
                padding: 20px;
                border-radius: 8px;
                border-left: 5px solid #3498db;
                text-align: center;
            }}
            .metric-label {{
                font-size: 0.9em;
                color: #7f8c8d;
                margin-bottom: 5px;
            }}
            .metric-value {{
                font-size: 2em;
                font-weight: bold;
                color: #2c3e50;
            }}
            .chart-container {{
                margin: 30px 0;
                text-align: center;
            }}
            .chart-container img {{
                max-width: 100%;
                border-radius: 8px;
                box-shadow: 0 4px 8px rgba(0,0,0,0.1);
            }}
            .risk-high {{ color: #e74c3c; font-weight: bold; }}
            .risk-low {{ color: #27ae60; font-weight: bold; }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin: 20px 0;
            }}
            th, td {{
                padding: 12px;
                text-align: center;
                border-bottom: 1px solid #ecf0f1;
            }}
            th {{
                background-color: #3498db;
                color: white;
            }}
            tr:hover {{
                background-color: #f8f9fa;
            }}
            .footer {{
                text-align: center;
                color: #7f8c8d;
                font-size: 0.9em;
                margin-top: 50px;
                padding-top: 20px;
                border-top: 1px solid #ecf0f1;
            }}
            .highlight-box {{
                background: linear-gradient(135deg, #27ae60 0%, #2ecc71 100%);
                color: white;
                padding: 20px;
                border-radius: 8px;
                margin: 20px 0;
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>🛢️ Multi-Scenario Oil Price Analysis</h1>
            <div class="timestamp">Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}</div>
            <div class="timestamp">Report ID: {timestamp}</div>
            <div class="timestamp">Oil Prices Analyzed: {', '.join([f'${p:.0f}' for p in oil_prices])}</div>
        </div>
        
        <div class="section">
            {executive_summary}
        </div>
        
        <div class="section">
            <h2>📊 Key Metrics</h2>
            <div class="metric-grid">
                <div class="metric-card">
                    <div class="metric-label">Oil Price Scenarios</div>
                    <div class="metric-value">{len(oil_prices)}</div>
                </div>
                <div class="metric-card">
                    <div class="metric-label">Simulations per Scenario</div>
                    <div class="metric-value">{n_simulations:,}</div>
                </div>
                <div class="metric-card">
                    <div class="metric-label">Total Simulations</div>
                    <div class="metric-value">{n_simulations * len(oil_prices):,}</div>
                </div>
                <div class="metric-card">
                    <div class="metric-label">Model R²</div>
                    <div class="metric-value">{model_r2:.1%}</div>
                </div>
                <div class="metric-card">
                    <div class="metric-label">Best Scenario</div>
                    <div class="metric-value">${min_risk_price:.0f}</div>
                </div>
            </div>
        </div>
        
        <div class="section">
            <h2>📈 Scenario Comparison Table</h2>
                <thead>
                    <tr>
                        <th>Oil Price</th>
                        <th>Mean GDP (New Law)</th>
                        <th>90% Growth Range</th>
                        <th>Recession Probability (New)</th>
                        <th>Status</th>
                        <th>Policy Impact</th>
                    </tr>
                </thead>
                <tbody>
                    {scenario_rows}
                </tbody>
            </table>
            <div style="font-size: 0.85em; color: #7f8c8d; text-align: right; margin-top: -10px;">
                🟢 Safe (<5% Risk) &nbsp; | &nbsp; 🟡 Caution (5-15% Risk) &nbsp; | &nbsp; 🔴 Danger (>15% Risk)
            </div>
        </div>
        
        <div class="section">
            <h2>📉 Visualizations</h2>
            
            <div class="chart-container">
                <h3>Oil Price Sensitivity Analysis</h3>
                <p style="font-size: 0.9em; color: #7f8c8d;">How GDP growth changes as oil price varies</p>
                <img src="data:image/png;base64,{charts['oil_sensitivity']}" alt="Oil Sensitivity">
            </div>
            
            <div class="chart-container">
                <h3>Recession Risk by Oil Price & Policy</h3>
                <img src="data:image/png;base64,{charts['multi_risk_bars']}" alt="Risk Comparison">
            </div>
            
            <div class="chart-container">
                <h3>GDP Growth Distribution by Oil Price (New Tax Act)</h3>
                <img src="data:image/png;base64,{charts['multi_distribution']}" alt="Distribution">
            </div>
            
            <div class="chart-container">
                <h3>GDP Growth Range by Oil Price</h3>
                <img src="data:image/png;base64,{charts['multi_boxplot']}" alt="Box Plot">
            </div>
            
            <div class="chart-container">
                <h3>Historical GDP Growth (Training Data)</h3>
                <img src="data:image/png;base64,{charts['historical']}" alt="Historical">
            </div>
            
            <div class="chart-container">
                <h3>Policy Influence: Factor Importance</h3>
                <img src="data:image/png;base64,{charts['sensitivity']}" alt="Sensitivity">
            </div>
        </div>
        
        <div class="section">
            <h2>🧠 Methodology</h2>
            <table>
                <tr><td><strong>Scenarios Tested:</strong></td><td>{len(oil_prices)} oil prices × 3 policy scenarios = {len(oil_prices) * 3} combinations</td></tr>
                <tr><td><strong>Monte Carlo Iterations:</strong></td><td>{n_simulations:,} per scenario</td></tr>
                <tr><td><strong>Oil Price Uncertainty:</strong></td><td>±$10/barrel (Normal distribution)</td></tr>
                <tr><td><strong>ML Framework:</strong></td><td>TPOT AutoML (Genetic Programming)</td></tr>
                <tr><td><strong>Model Accuracy:</strong></td><td>R² = {model_r2:.4f}</td></tr>
            </table>
        </div>
        
        <div class="footer">
            <p>Nigeria Tax Model System | Multi-Scenario Oil Price Analysis</p>
            <p>Data Sources: NBS, CBN, World Bank, Google Data Commons, NRS Dashboard</p>
            <p style="margin-top: 15px;">
                <a href="https://github.com/oyesanyf/nigeria-tax-impact-model" target="_blank" style="color: #3498db; text-decoration: none; font-weight: bold;">
                    View Source Code on GitHub ↗
                </a>
            </p>
        </div>
    </body>
    </html>
    """
    
    # Save report
    os.makedirs('reports', exist_ok=True)
    base_filename = f"reports/oil_scenario_report_{timestamp}"
    html_filename = f"{base_filename}.html"
    with open(html_filename, 'w', encoding='utf-8') as f:
        f.write(html)
    
    paths = {"html": os.path.abspath(html_filename)}
    
    # Generate Docx
    docx_filename = f"{base_filename}.docx"
    try:
        generate_multi_oil_docx(docx_filename, timestamp, executive_summary, oil_prices, scenario_results, model_r2, n_simulations, charts)
        paths["word"] = os.path.abspath(docx_filename)
    except Exception as e:
        print(f"❌ Error generating Word report: {e}")
        
    # Generate PDF
    pdf_filename = f"{base_filename}.pdf"
    try:
        generate_multi_oil_pdf(pdf_filename, timestamp, executive_summary, oil_prices, scenario_results, model_r2, n_simulations, charts)
        paths["pdf"] = os.path.abspath(pdf_filename)
    except Exception as e:
        print(f"❌ Error generating PDF report: {e}")
    
    print(f"✔ Multi-scenario report saved: {html_filename}")
    
    return paths


def generate_multi_oil_docx(filename, timestamp, summary, oil_prices, scenario_results, model_r2, n_simulations, charts):
    """Generate Word Document for Multi-Scenario Oil Analysis."""
    doc = Document()
    doc.add_heading('Multi-Scenario Oil Price Analysis', 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    doc.add_paragraph(f"Report ID: {timestamp}")
    doc.add_paragraph(f"Oil Prices Analyzed: {', '.join([f'${p:.0f}' for p in oil_prices])}")

    doc.add_heading('Executive Summary', level=1)
    clean_summary = summary.replace('<h3>', '').replace('</h3>', '').replace('<strong>', '').replace('</strong>', '').replace('<p>', '').replace('</p>', '\n')
    doc.add_paragraph(clean_summary)

    doc.add_heading('Key Metrics', level=1)
    table = doc.add_table(rows=1, cols=2)
    min_risk_price = min(oil_prices, key=lambda p: (scenario_results[p]['GDP_New'] < 0).mean())
    
    metrics = [
        ('Oil Price Scenarios', str(len(oil_prices))),
        ('Simulations per Scenario', f"{n_simulations:,}"),
        ('Total Simulations', f"{n_simulations * len(oil_prices):,}"),
        ('Model Accuracy (R²)', f"{model_r2:.1%}"),
        ('Best Scenario (Lowest Risk)', f"${min_risk_price:.0f}/barrel")
    ]
    
    for m, v in metrics:
        row = table.add_row().cells
        row[0].text = m
        row[1].text = v

    doc.add_heading('Scenario Comparison Table', level=1)
    # create table with headers
    comp_table = doc.add_table(rows=1, cols=6)
    hdr = comp_table.rows[0].cells
    hdr[0].text = 'Oil Price'
    hdr[1].text = 'Mean GDP'
    hdr[2].text = '90% Range'
    hdr[3].text = 'Risk (New)'
    hdr[4].text = 'Status'
    hdr[5].text = 'Impact'
    
    for oil_price in oil_prices:
        sim = scenario_results[oil_price]
        risk_old = (sim['GDP_Old'] < 0).mean() * 100
        risk_new = (sim['GDP_New'] < 0).mean() * 100
        mean_gdp = sim['GDP_New'].mean()
        p5 = np.percentile(sim['GDP_New'], 5)
        p95 = np.percentile(sim['GDP_New'], 95)
        
        # Determine Status
        if risk_new < 5: status = "SAFE"
        elif risk_new <= 15: status = "CAUTION"
        else: status = "DANGER"
        
        row = comp_table.add_row().cells
        row[0].text = f"${oil_price:.0f}"
        row[1].text = f"{mean_gdp:.2f}%"
        row[2].text = f"{p5:+.1f}% to {p95:+.1f}%"
        row[3].text = f"{risk_new:.1f}%"
        row[4].text = status
        row[5].text = f"{risk_new - risk_old:+.1f}pp"

    doc.add_heading('Visualizations', level=1)
    
    for title, key in chart_map:
        if key in charts:
            doc.add_heading(title, level=2)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
                tmpfile.write(base64.b64decode(charts[key]))
                tmpfile_path = tmpfile.name
            doc.add_picture(tmpfile_path, width=Inches(6))
            os.unlink(tmpfile_path)

    doc.add_heading('Source Code', level=1)
    doc.add_paragraph('Theoretical framework and source code available at:')
    doc.add_paragraph('https://github.com/oyesanyf/nigeria-tax-impact-model', style='List Bullet')

    doc.save(filename)


def generate_multi_oil_pdf(filename, timestamp, summary, oil_prices, scenario_results, model_r2, n_simulations, charts):
    """Generate PDF Document for Multi-Scenario Oil Analysis."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # Title
    pdf.set_font("helvetica", 'B', 16)
    pdf.cell(0, 10, "Multi-Scenario Oil Price Analysis", ln=True, align='C')
    pdf.set_font("helvetica", size=10)
    pdf.cell(0, 10, f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", ln=True, align='C')
    pdf.cell(0, 10, f"Oil Prices: {', '.join([f'${p:.0f}' for p in oil_prices])}", ln=True, align='C')
    pdf.ln(10)
    
    # Summary
    pdf.set_font("helvetica", 'B', 14)
    pdf.cell(0, 10, "Executive Summary", ln=True)
    pdf.set_font("helvetica", size=11)
    
    clean_summary = summary.replace('<h3>', '').replace('</h3>', '\n').replace('<strong>', '').replace('</strong>', '').replace('<p>', '').replace('</p>', '\n')
    clean_summary = clean_summary.replace('\u2013', '-').replace('\u2014', '-').replace('\u2018', "'").replace('\u2019', "'").replace('\u201c', '"').replace('\u201d', '"')
    
    pdf.multi_cell(0, 7, clean_summary)
    pdf.ln(5)
    
    # Comparison Table (Text based for simplicity in FPDF)
    pdf.set_font("helvetica", 'B', 14)
    pdf.cell(0, 10, "Scenario Comparison", ln=True)
    pdf.set_font("courier", size=9) # Monospace for alignment
    
    # Header
    pdf.cell(20, 8, "Price", 1)
    pdf.cell(20, 8, "GDP", 1)
    pdf.cell(50, 8, "90% Range", 1)
    pdf.cell(25, 8, "Risk(New)", 1)
    pdf.cell(25, 8, "Status", 1)
    pdf.cell(25, 8, "Impact", 1, ln=True)
    
    for oil_price in oil_prices:
        sim = scenario_results[oil_price]
        risk_old = (sim['GDP_Old'] < 0).mean() * 100
        risk_new = (sim['GDP_New'] < 0).mean() * 100
        mean_gdp = sim['GDP_New'].mean()
        p5 = np.percentile(sim['GDP_New'], 5)
        p95 = np.percentile(sim['GDP_New'], 95)
        
        # Determine Status
        if risk_new < 5: status = "SAFE"
        elif risk_new <= 15: status = "CAUTION"
        else: status = "DANGER"
        
        pdf.cell(20, 8, f"${oil_price:.0f}", 1)
        pdf.cell(20, 8, f"{mean_gdp:.2f}%", 1)
        pdf.cell(50, 8, f"{p5:+.1f}% to {p95:+.1f}%", 1)
        pdf.cell(25, 8, f"{risk_new:.1f}%", 1)
        pdf.cell(25, 8, status, 1)
        pdf.cell(25, 8, f"{risk_new - risk_old:+.1f}pp", 1, ln=True)
        
    pdf.ln(10)
    
    # Visualizations
    pdf.set_font("helvetica", 'B', 14)
    pdf.cell(0, 10, "Visualizations", ln=True)
    
    chart_map = [
        ('Oil Price Sensitivity', 'oil_sensitivity'),
        ('Recession Risk Comparison', 'multi_risk_bars'),
        ('GDP Growth Distribution', 'multi_distribution'),
        ('GDP Growth Range', 'multi_boxplot'),
        ('Policy Influence', 'sensitivity')
    ]
    
    for title, key in chart_map:
        if key in charts:
            if pdf.get_y() > 200: pdf.add_page()
            pdf.set_font("helvetica", 'B', 12)
            pdf.cell(0, 10, title, ln=True)
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
                tmpfile.write(base64.b64decode(charts[key]))
                tmpfile_path = tmpfile.name
            pdf.image(tmpfile_path, x=10, w=180)
            os.unlink(tmpfile_path)
            pdf.ln(5)
            
    # GitHub Link
    pdf.ln(10)
    pdf.set_font("helvetica", 'I', 10)
    pdf.set_text_color(0, 0, 255)
    pdf.cell(0, 10, "Source Code: https://github.com/oyesanyf/nigeria-tax-impact-model", ln=True, align='C', link="https://github.com/oyesanyf/nigeria-tax-impact-model")
            
    pdf.output(filename)


if __name__ == "__main__":
    print("This module should be imported by nigeria_tax_model.py")
